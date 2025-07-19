# main.py - Updated with real AI processing
import os
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any
import asyncio

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# Document processing imports
import PyPDF2
from PIL import Image
import pytesseract
from docx import Document as DocxDocument

# AI/LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.llms import OpenAI
from langchain.schema import Document
import openai

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Initialize FastAPI
app = FastAPI(title="Document Research API")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Pydantic models
class QueryRequest(BaseModel):
    question: str

# Global variables
uploaded_files = []
processed_documents = []
vectorstore = None
qa_chain = None

# Document processing functions
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""

def extract_text_from_image(file_path: str) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"Error extracting image text: {e}")
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        print(f"Error reading text file: {e}")
        return ""

def process_document(file_path: str, filename: str) -> str:
    """Process document and extract text based on file type"""
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        return extract_text_from_docx(file_path)
    elif file_ext in ['.jpg', '.jpeg', '.png']:
        return extract_text_from_image(file_path)
    elif file_ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        return ""

# API Routes
@app.get("/")
@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "message": "Document Research API is running",
        "timestamp": datetime.now().isoformat(),
        "ai_enabled": bool(os.getenv("OPENAI_API_KEY"))
    }

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Save file
        file_path = UPLOAD_DIR / file.filename
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Track uploaded file
        uploaded_files.append({
            "filename": file.filename,
            "path": str(file_path),
            "timestamp": datetime.now().isoformat()
        })
        
        return {
            "message": f"File {file.filename} uploaded successfully",
            "total_files": len(uploaded_files)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/ingest/")
async def ingest_documents():
    global vectorstore, qa_chain, processed_documents
    
    try:
        if not uploaded_files:
            raise HTTPException(status_code=400, detail="No files uploaded")
        
        if not os.getenv("OPENAI_API_KEY"):
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # Process each uploaded file
        processed_documents = []
        all_texts = []
        
        for file_info in uploaded_files:
            file_path = file_info["path"]
            filename = file_info["filename"]
            
            # Extract text from document
            text_content = process_document(file_path, filename)
            
            if text_content:
                processed_documents.append({
                    "filename": filename,
                    "content": text_content,
                    "timestamp": file_info["timestamp"]
                })
                
                # Create LangChain document
                doc = Document(
                    page_content=text_content,
                    metadata={"source": filename, "timestamp": file_info["timestamp"]}
                )
                all_texts.append(doc)
        
        if not all_texts:
            raise HTTPException(status_code=400, detail="No text content extracted from documents")
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        splits = text_splitter.split_documents(all_texts)
        
        # Create embeddings and vector store
        embeddings = OpenAIEmbeddings()
        vectorstore = Chroma.from_documents(splits, embeddings)
        
        # Create QA chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=OpenAI(temperature=0),
            chain_type="stuff",
            retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
            return_source_documents=True
        )
        
        return {
            "message": "Documents processed successfully with AI",
            "processed_files": len(processed_documents),
            "total_chunks": len(splits),
            "status": "ready"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@app.post("/query/")
async def query_documents(request: QueryRequest):
    try:
        if not qa_chain:
            raise HTTPException(status_code=400, detail="Documents not processed yet. Please ingest documents first.")
        
        if not request.question.strip():
            raise HTTPException(status_code=400, detail="Question cannot be empty")
        
        # Get AI-powered answer
        result = qa_chain({"query": request.question})
        
        # Extract sources
        sources = []
        if "source_documents" in result:
            sources = [doc.metadata.get("source", "Unknown") for doc in result["source_documents"]]
        
        return {
            "answer": result["result"],
            "question": request.question,
            "sources": list(set(sources)),  # Remove duplicates
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query failed: {str(e)}")

@app.get("/status/")
async def get_status():
    return {
        "uploaded_files": len(uploaded_files),
        "processed_documents": len(processed_documents),
        "vectorstore_ready": vectorstore is not None,
        "qa_chain_ready": qa_chain is not None,
        "ai_configured": bool(os.getenv("OPENAI_API_KEY"))
    }

@app.get("/documents/")
async def list_documents():
    """List processed documents with preview"""
    return {
        "documents": [
            {
                "filename": doc["filename"],
                "preview": doc["content"][:200] + "..." if len(doc["content"]) > 200 else doc["content"],
                "length": len(doc["content"]),
                "timestamp": doc["timestamp"]
            }
            for doc in processed_documents
        ]
    }

@app.delete("/clear/")
async def clear_data():
    global vectorstore, qa_chain, uploaded_files, processed_documents
    
    try:
        # Clear files
        for file_info in uploaded_files:
            file_path = Path(file_info["path"])
            if file_path.exists():
                file_path.unlink()
        
        # Reset all variables
        uploaded_files = []
        processed_documents = []
        vectorstore = None
        qa_chain = None
        
        return {"message": "All data cleared successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Clear failed: {str(e)}")

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)